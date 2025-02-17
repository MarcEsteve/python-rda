import docx # pip install python-docx

ruta= 'C:/Users/artha/OneDrive/Escritorio/AKKODIS/python-rda/05_xlsx_docx_pdf/tecnologias_emergentes.docx'

#Lectura
d = docx.Document(ruta) #Cargamos el documento docx
d.paragraphs #Lista con los párrafos del documento

d.paragraphs[0].text #Texto del primer párrafo

p= d.paragraphs[2] #Guardamos el texto del tercer párrafo en la variable p

p.runs[0].text #Texto del primer run del párrafo p

p.runs[1].text #Texto del segundo run del párrafo p

p.runs[2].text #Texto del tercer run del párrafo p

#Lectura completa
def coger_texto(fichero):
    doc= docx.Document(fichero)
    texto_completo = []
    for parrafo in doc.paragraphs:
        texto_completo.append(parrafo.text)
    return "\n".join(texto_completo)

print(coger_texto(ruta))

# Coger texto sin líneas vacías
def coger_texto(fichero):
    doc = docx.Document(fichero)
    texto_completo = []
    
    for parrafo in doc.paragraphs:
        texto = parrafo.text.strip()  # Eliminar espacios y saltos de línea al inicio y al final
        
        if texto:  # Solo agregamos el párrafo si no está vacío
            texto_completo.append(texto)
    
    return "\n\n".join(texto_completo)  # Usar doble salto de línea para mejor separación

#Edición

p.runs[1].bold = True #Negrita, para cursiva seria italic y underline para subrayado, por ejemplo

p.runs[2].text = 'otro texto' #Cambiamos el texto del segundo run del párrafo p

d.save(r'C:/Users/artha/OneDrive/Escritorio/AKKODIS/python-rda/05_xlsx_docx_pdf/tecnologias_emergentes_v2.docx') #Guardamos el documento